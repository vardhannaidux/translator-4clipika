# -*- coding: utf-8 -*-
"""
services.py — Core translation processing pipelines for text and document formats (TXT, DOCX).
"""

import sys
import os
import time
from io import BytesIO
from typing import Tuple, Dict, Any

# Ensure parent directory is in search path to access root modules
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), "..")))

from translation_engine import translate_text, de_render
from config import setup_logger

logger = setup_logger("backend_services")

def process_text_translation(text: str, direction: str, editorial_mode: bool) -> Tuple[str, float]:
    """
    Translates a plain text string in the specified direction.
    
    Args:
        text (str): Input text string.
        direction (str): "unicode_to_legacy" or "legacy_to_unicode".
        editorial_mode (bool): Toggles spelling corrections in forward mode.
        
    Returns:
        Tuple[str, float]: The translated string and the execution time in ms.
    """
    start_time = time.perf_counter()
    logger.info(f"Processing text translation of length {len(text)} in direction '{direction}'")
    
    if direction == "unicode_to_legacy":
        result = translate_text(text, editorial_mode=editorial_mode)
    elif direction == "legacy_to_unicode":
        # Legacy input text is passed as a string representing CP1252 characters.
        # We must encode it to bytes, then pass it to the de_render decoder.
        raw_bytes = text.encode("cp1252", errors="replace")
        result = de_render(raw_bytes)
    else:
        raise ValueError(f"Invalid translation direction: '{direction}'")
        
    elapsed_ms = (time.perf_counter() - start_time) * 1000.0
    logger.info(f"Text translation completed in {elapsed_ms:.2f} ms")
    return result, elapsed_ms


def translate_txt_file(file_bytes: bytes, direction: str, editorial_mode: bool) -> bytes:
    """
    Translates a plain text file.
    """
    # 1. Decode inputs based on direction
    if direction == "unicode_to_legacy":
        input_text = file_bytes.decode("utf-8", errors="replace")
        translated, _ = process_text_translation(input_text, direction, editorial_mode)
        # Output is legacy bytes (CP1252)
        return translated.encode("cp1252", errors="replace")
    else:
        input_text = file_bytes.decode("cp1252", errors="replace")
        translated, _ = process_text_translation(input_text, direction, editorial_mode)
        # Output is Unicode (UTF-8)
        return translated.encode("utf-8", errors="replace")


def translate_docx_file(file_bytes: bytes, direction: str, editorial_mode: bool) -> bytes:
    """
    Translates a Word DOCX file run-by-run to preserve layout and typography.
    """
    import docx
    
    logger.info("Starting DOCX document translation")
    doc = docx.Document(BytesIO(file_bytes))
    
    def process_str(text: str) -> str:
        if not text.strip():
            return text
        if direction == "unicode_to_legacy":
            return translate_text(text, editorial_mode=editorial_mode)
        else:
            raw = text.encode("cp1252", errors="replace")
            return de_render(raw)
            
    # Convert body paragraphs
    for p in doc.paragraphs:
        for run in p.runs:
            if run.text.strip():
                run.text = process_str(run.text)
                
    # Convert table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        if run.text.strip():
                            run.text = process_str(run.text)
                            
    out_stream = BytesIO()
    doc.save(out_stream)
    logger.info("DOCX document translation completed successfully")
    return out_stream.getvalue()
